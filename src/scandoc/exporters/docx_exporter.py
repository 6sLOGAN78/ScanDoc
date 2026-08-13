"""
DocxExporter serializing DocumentIR into Microsoft Word DOCX binary documents.
"""

import io
from typing import List, Optional

import docx

from scandoc.exporters.base import BaseExporter
from scandoc.exporters.models import ExportOptions, ExportResult
from scandoc.exporters.taxonomy import OutputDestination
from scandoc.models import DocumentIR
from scandoc.models.blocks import (
    FigureBlock,
    FormulaBlock,
    HeadingBlock,
    ListBlock,
    ParagraphBlock,
    TableBlock,
)


class DocxExporter(BaseExporter):
    """
    Exporter converting DocumentIR into Microsoft Word (.docx) documents.
    """

    @property
    def format_id(self) -> str:
        return "docx"

    @property
    def description(self) -> str:
        return "Microsoft Word DOCX Exporter"

    def export(
        self,
        document: DocumentIR,
        options: Optional[ExportOptions] = None,
    ) -> ExportResult:
        opts = options or ExportOptions(format_id="docx", destination=OutputDestination.BYTES)
        doc = docx.Document()
        warnings: List[str] = []
        asset_refs: List[str] = []

        if opts.include_metadata and document.metadata:
            doc.add_heading(document.metadata.name or "Document", level=0)
            if document.metadata.author:
                doc.add_paragraph(f"Author: {document.metadata.author}")

        for page in document.pages:
            for block in page.blocks:
                if isinstance(block, HeadingBlock):
                    lvl = max(1, min(4, block.level))
                    doc.add_heading(block.text, level=lvl)

                elif isinstance(block, ParagraphBlock):
                    doc.add_paragraph(block.text)

                elif isinstance(block, ListBlock):
                    is_ord = getattr(block, "ordered", getattr(block, "is_ordered", False))
                    style = "List Number" if is_ord else "List Bullet"
                    for item in block.items:
                        item_txt = getattr(item, "text", str(item))
                        doc.add_paragraph(item_txt, style=style)

                elif isinstance(block, TableBlock):
                    grid = [["" for _ in range(block.num_cols)] for _ in range(block.num_rows)]
                    for cell in block.cells:
                        if cell.row_index < block.num_rows and cell.col_index < block.num_cols:
                            grid[cell.row_index][cell.col_index] = cell.text or ""
                    if grid:
                        tbl = doc.add_table(rows=len(grid), cols=len(grid[0]))
                        for r_idx, row in enumerate(grid):
                            for c_idx, cell_text in enumerate(row):
                                tbl.cell(r_idx, c_idx).text = cell_text or ""

                elif isinstance(block, FigureBlock):
                    img_bytes = None
                    if hasattr(block, "image_ref") and block.image_ref and block.image_ref.base64_data:
                        import base64
                        img_bytes = base64.b64decode(block.image_ref.base64_data)
                    else:
                        img_bytes = getattr(block, "image_bytes", None)

                    if img_bytes:
                        try:
                            stream = io.BytesIO(img_bytes)
                            doc.add_picture(stream)
                        except Exception as e:
                            warnings.append(f"Failed to render image in DOCX: {e}")
                    if block.caption:
                        doc.add_paragraph(f"Caption: {block.caption}")

                elif isinstance(block, FormulaBlock):
                    doc.add_paragraph(f"Equation: {block.expression}")

                else:
                    text = getattr(block, "text", str(block))
                    doc.add_paragraph(text)

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()

        output_path = opts.output_path
        if opts.destination == OutputDestination.FILE_PATH and output_path:
            with open(output_path, "wb") as f:
                f.write(docx_bytes)

        return ExportResult(
            format_id="docx",
            destination=opts.destination,
            content=docx_bytes,
            output_path=output_path,
            warnings=warnings,
            asset_references=asset_refs,
        )
